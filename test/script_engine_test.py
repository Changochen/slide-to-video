import pytest
from slide_to_video.script_engine import ScriptEngine, Script, ScriptConfig
from slide_to_video.script_engine import extract_text_from_docx


@pytest.fixture
def script_engine():
    return ScriptEngine()


def test_extract_text_from_docx(tmp_path):
    from docx import Document

    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    doc.save(doc_path)
    result = extract_text_from_docx(doc_path)
    assert result == "Hello\nWorld"


def test_load_script_txt(script_engine, tmp_path):
    script_path = tmp_path / "test.txt"
    script_path.write_text("Hello World")
    result = script_engine.load_script(script_path)
    assert result == "Hello World"


def test_load_script_docx(script_engine, tmp_path):
    from docx import Document

    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.save(doc_path)
    result = script_engine.load_script(doc_path)
    assert result == "Hello World"


def test_parse_script_no_config(script_engine):
    script = "This is a script without config"
    result_text, result_config = script_engine.parse_script(script)
    assert result_text == "This is a script without config"
    assert result_config is None


def test_parse_script_with_config(script_engine):
    script = "This is a script with config===#delay: 2.5"
    result_text, result_config = script_engine.parse_script(script)
    assert result_text == "This is a script with config"
    assert result_config == ScriptConfig({"delay": 2.5})


def test_split_script(script_engine, tmp_path):
    script_path = tmp_path / "test.txt"
    script_path.write_text("Part1NEWSLIDEPart2")
    output_path = tmp_path / "output"
    output_path.mkdir()
    result = script_engine.split_script(script_path, str(output_path))
    assert len(result) == 2
    assert result[0].text == "Part1"
    assert result[1].text == "Part2"
